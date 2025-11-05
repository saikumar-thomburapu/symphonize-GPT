"""
File Upload and Processing Service - With Image Base64 Support
"""

import io
import logging
import base64
from typing import Dict, Any

logger = logging.getLogger(__name__)

# File size limit (10MB)
MAX_FILE_SIZE = 10 * 1024 * 1024

# Allowed mime types
ALLOWED_TYPES = {
    'application/pdf': 'pdf',
    'text/plain': 'txt',
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
    'image/png': 'png',
    'image/jpeg': 'jpg',
    'image/jpg': 'jpg',
    'image/webp': 'webp',
}

class FileService:
    """Service for handling file uploads and text extraction."""
    
    def validate_file(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Validate uploaded file."""
        if len(file_content) > MAX_FILE_SIZE:
            raise Exception(f"File size exceeds {MAX_FILE_SIZE / 1024 / 1024}MB limit")
        
        if mime_type not in ALLOWED_TYPES:
            raise Exception(f"File type {mime_type} not allowed")
        
        return {
            'filename': filename,
            'size': len(file_content),
            'mime_type': mime_type,
            'file_type': ALLOWED_TYPES[mime_type],
        }
    
    async def extract_text_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += f"\n--- Page {page_num + 1} ---\n"
                text += page.extract_text()
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise Exception(f"Error processing PDF: {str(e)}")
    
    async def extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text = ""
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    async def extract_text_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file."""
        try:
            return file_content.decode('utf-8').strip()
        except Exception as e:
            logger.error(f"Error reading TXT file: {str(e)}")
            raise Exception(f"Error processing TXT: {str(e)}")
    
    async def convert_image_to_base64(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, str]:
        """Convert image to base64 for Qwen3 vision model."""
        try:
            # Validate image
            from PIL import Image
            img = Image.open(io.BytesIO(file_content))
            
            # Get media type
            media_type_map = {
                'image/png': 'image/png',
                'image/jpeg': 'image/jpeg',
                'image/jpg': 'image/jpeg',
                'image/webp': 'image/webp',
            }
            media_type = media_type_map.get(mime_type, 'image/jpeg')
            
            # Convert to base64
            base64_image = base64.b64encode(file_content).decode('utf-8')
            
            logger.info(f"✓ Image converted to base64: {filename} ({len(base64_image)} bytes)")
            
            return {
                'type': 'image',
                'media_type': media_type,
                'data': base64_image,
                'filename': filename,
            }
        except Exception as e:
            logger.error(f"Error converting image to base64: {str(e)}")
            raise Exception(f"Error processing image: {str(e)}")
    
    async def process_file(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """Process uploaded file and extract content."""
        file_info = self.validate_file(file_content, filename, mime_type)
        file_type = file_info['file_type']
        
        logger.info(f"Processing file: {filename} ({file_type})")
        
        # For images, return base64 data
        if file_type in ['jpg', 'png', 'webp']:
            logger.info(f"📸 Image detected: {filename} - Converting to base64")
            image_data = await self.convert_image_to_base64(file_content, filename, mime_type)
            return {
                'filename': filename,
                'file_type': file_type,
                'is_image': True,
                'image_data': image_data,
                'extracted_text': f"[Image: {filename}]",  # Placeholder for text
            }
        
        # For text files, extract text
        extracted_text = ""
        
        if file_type == 'pdf':
            logger.info(f"📄 PDF detected: {filename} - Extracting text")
            extracted_text = await self.extract_text_from_pdf(file_content)
        elif file_type == 'docx':
            logger.info(f"📝 DOCX detected: {filename} - Extracting text")
            extracted_text = await self.extract_text_from_docx(file_content)
        elif file_type == 'txt':
            logger.info(f"📋 TXT detected: {filename} - Reading text")
            extracted_text = await self.extract_text_from_txt(file_content)
        
        return {
            'filename': filename,
            'file_type': file_type,
            'is_image': False,
            'extracted_text': extracted_text,
        }

# Singleton instance
file_service = FileService()











# """
# File Upload and Processing Service
# """

# import io
# import logging
# from typing import Dict, Any

# logger = logging.getLogger(__name__)

# # File size limit (10MB)
# MAX_FILE_SIZE = 10 * 1024 * 1024

# # Allowed mime types
# ALLOWED_TYPES = {
#     'application/pdf': 'pdf',
#     'text/plain': 'txt',
#     'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
#     'image/png': 'png',
#     'image/jpeg': 'jpg',
#     'image/jpg': 'jpg',
#     'image/webp': 'webp',
# }


# class FileService:
#     """Service for handling file uploads and text extraction."""
    
#     def validate_file(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
#         """Validate uploaded file."""
#         if len(file_content) > MAX_FILE_SIZE:
#             raise Exception(f"File size exceeds {MAX_FILE_SIZE / 1024 / 1024}MB limit")
        
#         if mime_type not in ALLOWED_TYPES:
#             raise Exception(f"File type {mime_type} not allowed")
        
#         return {
#             'filename': filename,
#             'size': len(file_content),
#             'mime_type': mime_type,
#             'file_type': ALLOWED_TYPES[mime_type],
#         }
    
#     async def extract_text_from_pdf(self, file_content: bytes) -> str:
#         """Extract text from PDF file."""
#         try:
#             import PyPDF2
#             pdf_file = io.BytesIO(file_content)
#             pdf_reader = PyPDF2.PdfReader(pdf_file)
            
#             text = ""
#             for page_num in range(len(pdf_reader.pages)):
#                 page = pdf_reader.pages[page_num]
#                 text += f"\n--- Page {page_num + 1} ---\n"
#                 text += page.extract_text()
            
#             return text.strip()
#         except Exception as e:
#             logger.error(f"Error extracting PDF text: {str(e)}")
#             raise Exception(f"Error processing PDF: {str(e)}")
    
#     async def extract_text_from_docx(self, file_content: bytes) -> str:
#         """Extract text from DOCX file."""
#         try:
#             from docx import Document
#             docx_file = io.BytesIO(file_content)
#             doc = Document(docx_file)
            
#             text = ""
#             for para in doc.paragraphs:
#                 if para.text.strip():
#                     text += para.text + "\n"
            
#             return text.strip()
#         except Exception as e:
#             logger.error(f"Error extracting DOCX text: {str(e)}")
#             raise Exception(f"Error processing DOCX: {str(e)}")
    
#     async def extract_text_from_txt(self, file_content: bytes) -> str:
#         """Extract text from TXT file."""
#         try:
#             return file_content.decode('utf-8').strip()
#         except Exception as e:
#             logger.error(f"Error reading TXT file: {str(e)}")
#             raise Exception(f"Error processing TXT: {str(e)}")
    
#     async def extract_text_from_image(self, file_content: bytes, filename: str) -> str:
#         """Placeholder for image text extraction."""
#         return f"[Image: {filename} - Image analysis not yet implemented]"
    
#     async def process_file(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
#         """Process uploaded file and extract content."""
#         file_info = self.validate_file(file_content, filename, mime_type)
#         file_type = file_info['file_type']
        
#         logger.info(f"Processing file: {filename} ({file_type})")
        
#         extracted_text = ""
        
#         if file_type == 'pdf':
#             extracted_text = await self.extract_text_from_pdf(file_content)
#         elif file_type == 'docx':
#             extracted_text = await self.extract_text_from_docx(file_content)
#         elif file_type == 'txt':
#             extracted_text = await self.extract_text_from_txt(file_content)
#         elif file_type in ['jpg', 'png', 'webp']:
#             extracted_text = await self.extract_text_from_image(file_content, filename)
        
#         return {
#             'filename': filename,
#             'file_type': file_type,
#             'extracted_text': extracted_text,
#         }


# # Singleton instance
# file_service = FileService()
